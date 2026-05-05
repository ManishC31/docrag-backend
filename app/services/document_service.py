import logging
import os
import uuid
from pathlib import Path

logger = logging.getLogger(__name__)

import aiofiles
from fastapi import BackgroundTasks, HTTPException, UploadFile, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.config import settings
from app.core.constants import DocumentStatus, ErrorMessages, UserRole
from app.models.document import Document, DocumentChunk
from app.models.group import Group
from app.models.user import User
from app.schemas.document import DocumentResponse
from app.services.rag_service import process_document_chunks


async def _get_group_or_404(group_id: uuid.UUID, db: AsyncSession) -> Group:
    result = await db.execute(
        select(Group).where(Group.id == group_id).options(selectinload(Group.documents))
    )
    group = result.scalar_one_or_none()
    if not group:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ErrorMessages.GROUP_NOT_FOUND)
    return group


def _extract_text(file_path: str, file_type: str) -> str:
    if file_type == "txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    if file_type == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if file_type == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join(para.text for para in doc.paragraphs)

    raise ValueError(f"Unsupported file type: {file_type}")


def _chunk_text(text: str) -> list[str]:
    chunks = []
    start = 0
    size = settings.CHUNK_SIZE
    overlap = settings.CHUNK_OVERLAP

    while start < len(text):
        end = start + size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap

    return chunks


async def _process_document_background(document_id: str, file_path: str, file_type: str) -> None:
    from app.db.session import AsyncSessionLocal

    async with AsyncSessionLocal() as db:
        try:
            result = await db.execute(select(Document).where(Document.id == uuid.UUID(document_id)))
            document = result.scalar_one_or_none()
            if not document:
                return

            document.status = DocumentStatus.PROCESSING
            await db.commit()

            raw_text = _extract_text(file_path, file_type)
            if not raw_text.strip():
                raise ValueError("No text could be extracted from the document")

            chunks = _chunk_text(raw_text)
            await process_document_chunks(document_id=uuid.UUID(document_id), chunks=chunks, db=db)

            document.status = DocumentStatus.READY
            await db.commit()
        except Exception as exc:
            logger.exception("Document processing failed for %s: %s", document_id, exc)
            await db.rollback()
            async with AsyncSessionLocal() as err_db:
                result = await err_db.execute(select(Document).where(Document.id == uuid.UUID(document_id)))
                doc = result.scalar_one_or_none()
                if doc:
                    doc.status = DocumentStatus.FAILED
                    doc.error_message = str(exc)[:500]
                    await err_db.commit()


async def upload_document(
    group_id: uuid.UUID,
    file: UploadFile,
    current_user: User,
    db: AsyncSession,
    background_tasks: BackgroundTasks,
) -> DocumentResponse:
    group = await _get_group_or_404(group_id, db)

    if current_user.role != UserRole.ADMIN and group.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=ErrorMessages.GROUP_ACCESS_DENIED)

    if len(group.documents) >= settings.MAX_DOCUMENTS_PER_GROUP:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=ErrorMessages.DOCUMENT_LIMIT_EXCEEDED)

    if not file.filename:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No filename provided")

    extension = file.filename.rsplit(".", 1)[-1].lower()
    if extension not in settings.ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=ErrorMessages.UNSUPPORTED_FILE_TYPE)

    content = await file.read()
    if len(content) > settings.MAX_FILE_SIZE_BYTES:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=ErrorMessages.FILE_TOO_LARGE)

    # Save file to disk
    upload_dir = Path(settings.UPLOAD_DIR) / str(current_user.id) / str(group_id)
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_filename = f"{uuid.uuid4()}.{extension}"
    file_path = upload_dir / safe_filename

    async with aiofiles.open(file_path, "wb") as f:
        await f.write(content)

    document = Document(
        name=file.filename.rsplit(".", 1)[0],
        original_filename=file.filename,
        file_type=extension,
        file_size=len(content),
        file_path=str(file_path),
        group_id=group_id,
        status=DocumentStatus.PENDING,
    )
    db.add(document)
    await db.flush()
    await db.commit()  # must commit before background task runs, which opens its own session

    background_tasks.add_task(
        _process_document_background,
        str(document.id),
        str(file_path),
        extension,
    )

    return DocumentResponse.model_validate(document)


async def list_documents(group_id: uuid.UUID, current_user: User, db: AsyncSession) -> list[DocumentResponse]:
    group = await _get_group_or_404(group_id, db)

    if current_user.role != UserRole.ADMIN and group.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=ErrorMessages.GROUP_ACCESS_DENIED)

    result = await db.execute(
        select(Document).where(Document.group_id == group_id).order_by(Document.created_at.desc())
    )
    documents = result.scalars().all()
    return [DocumentResponse.model_validate(d) for d in documents]


async def delete_document(document_id: uuid.UUID, current_user: User, db: AsyncSession) -> None:
    from qdrant_client.models import PointIdsList
    from app.db.qdrant import get_qdrant_client

    result = await db.execute(
        select(Document)
        .where(Document.id == document_id)
        .options(selectinload(Document.group), selectinload(Document.chunks))
    )
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ErrorMessages.DOCUMENT_NOT_FOUND)

    if current_user.role != UserRole.ADMIN and document.group.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=ErrorMessages.GROUP_ACCESS_DENIED)

    chunk_ids = [str(chunk.id) for chunk in document.chunks]
    if chunk_ids:
        qdrant = get_qdrant_client()
        await qdrant.delete(
            collection_name=settings.QDRANT_COLLECTION_NAME,
            points_selector=PointIdsList(points=chunk_ids),
        )

    try:
        os.remove(document.file_path)
    except FileNotFoundError:
        pass

    await db.delete(document)
